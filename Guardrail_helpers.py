import os, io, zipfile, base64, tempfile, json, datetime, re
import pandas as pd
import streamlit as st
from pathlib import Path
from typing import List, Union
from PIL import Image

# LangChain imports
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma

# OpenAI client
from openai import OpenAI

# ---------------------------
# Config
# ---------------------------
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 100
CHROMA_DIR = "./chroma_store"
EMBEDDING_MODEL = "text-embedding-3-small"  # keep consistent with existing DB (1536-d)

# ---------------------------
# (Temporary) Hardcoded API key - DO NOT COMMIT
# ---------------------------
# Replace the string below with your key for local testing only.
																													

# ---------------------------
# Guardrail helpers (defined before UI usage)
# ---------------------------
def moderate_text(text: str):
    """
    Run a moderation check on text. Returns (flagged: bool, reasons: List[str]).
    Robustly normalizes 'categories' from dict-like or attribute-based (pydantic) responses.
    """
    try:
        if not text or not text.strip():
            return False, []

        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        resp = client.moderations.create(model="omni-moderation-latest", input=text)

        # normalize to the first result object/dict
        if isinstance(resp, dict):
            res0 = resp.get("results", [resp])[0]
        else:
            res0 = getattr(resp, "results", [resp])[0]

        # extract flagged and raw categories
        if isinstance(res0, dict):
            flagged = bool(res0.get("flagged", False))
            categories_raw = res0.get("categories", {}) or {}
        else:
            flagged = bool(getattr(res0, "flagged", False))
            categories_raw = getattr(res0, "categories", {}) or {}

        # normalize categories into a plain dict of booleans
        def normalize_bool_map(obj):
            if obj is None:
                return {}
            if isinstance(obj, dict):
                return {k: bool(v) for k, v in obj.items()}
            if hasattr(obj, "dict") and callable(getattr(obj, "dict")):
                try:
                    return {k: bool(v) for k, v in obj.dict().items()}
                except Exception:
                    pass
            if hasattr(obj, "__dict__"):
                try:
                    return {k: bool(v) for k, v in vars(obj).items()}
                except Exception:
                    pass
            out = {}
            for k in dir(obj):
                if k.startswith("_"):
                    continue
                try:
                    v = getattr(obj, k)
                except Exception:
                    continue
                if isinstance(v, bool):
                    out[k] = v
            return out

        categories = normalize_bool_map(categories_raw)
        reasons = [k for k, v in categories.items() if v]

        return flagged, reasons

    except Exception as e:
        st.warning(f"Moderation check failed: {e}")
        return False, []

# ---------------------------
# Multimodal Embeddings 
# ---------------------------
class MultimodalEmbeddings:
    def __init__(self, openai_api_key: str = None):
        self.api_key = openai_api_key or os.environ["OPENAI_API_KEY"]
        self.client = OpenAI(api_key=self.api_key)

    def encode_image_to_base64(self, image_path: str) -> str:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode("utf-8")

    def get_image_description(self, image_path: str) -> str:
        base64_image = self.encode_image_to_base64(image_path)
        response = self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Describe this image in detail for compliance context."},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ],
                }
            ],
        )
        # Support both object-like and dict-like responses
        if isinstance(response, dict):
            return response["choices"][0]["message"]["content"]
        else:
            return response.choices[0].message.content

multi_embed = MultimodalEmbeddings()

# ---------------------------
# File extraction
# ---------------------------
def extract_text_from_file(file_name: str, file_bytes: bytes) -> List[Document]:
    _, ext = os.path.splitext(file_name.lower())
    docs: List[Document] = []
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(page_content=text, metadata={"source": file_name, "page": i+1}))
        elif ext in [".docx", ".doc"]:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([p.text for p in doc.paragraphs])
            if text.strip():
                docs.append(Document(page_content=text, metadata={"source": file_name}))
        elif ext in [".txt", ".csv"]:
            text = file_bytes.decode(errors="ignore")
            docs.append(Document(page_content=text, metadata={"source": file_name}))
        elif ext in [".png", ".jpg", ".jpeg"]:
            # Preview uploaded image in the Streamlit UI to help diagnose issues
            try:
                pil_preview = Image.open(io.BytesIO(file_bytes)).convert("RGB")
                st.image(pil_preview, caption=file_name, use_column_width=True)
            except Exception as e:
                st.warning(f"Could not open image preview: {e}")

            # Save temp image and validate contents (bbox None => fully empty/black)
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name

                pil = Image.open(tmp_path)
                if pil.getbbox() is None:
                    # all-black or empty image, record metadata rather than attempting vision call
                    docs.append(Document(page_content="", metadata={"source": file_name, "note": "empty_or_all_black_image"}))
                else:
                    # attempt to get a description (guarded)
                    try:
                        description = multi_embed.get_image_description(tmp_path)
                    except Exception as e:
                        st.warning(f"Image description call failed: {e}")
                        description = ""
                    docs.append(Document(page_content=description, metadata={"source": file_name, "type": "image_caption"}))
            except Exception as e:
                docs.append(Document(page_content="", metadata={"source": file_name, "error": str(e)}))
        elif ext == ".zip":
            with zipfile.ZipFile(io.BytesIO(file_bytes), 'r') as z:
                for name in z.namelist():
                    if name.endswith("/"):
                        continue
                    inner_bytes = z.read(name)
                    inner_docs = extract_text_from_file(name, inner_bytes)
                    for d in inner_docs:
                        d.metadata["source"] = f"{file_name} -> {d.metadata.get('source', name)}"
                        docs.append(d)
        else:
            docs.append(Document(page_content="", metadata={"source": file_name, "note": "unsupported_type"}))
    except Exception as e:
        docs.append(Document(page_content="", metadata={"source": file_name, "error": str(e)}))
    return docs

# ---------------------------
# Chunking
# ---------------------------
def chunk_documents(documents: List[Document], chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    all_chunks: List[Document] = []
    for doc in documents:
        chunks = splitter.split_text(doc.page_content or "")
        for i, c in enumerate(chunks):
            metadata = dict(doc.metadata) if doc.metadata else {}
            metadata.update({"chunk_index": i})
            all_chunks.append(Document(page_content=c, metadata=metadata))
    return all_chunks

# ---------------------------
# Helper: load persisted Chroma if present
# ---------------------------
def load_persisted_vectorstore():
    if not os.path.exists(CHROMA_DIR):
        return None
    embedding_model = OpenAIEmbeddings(model=EMBEDDING_MODEL)
    try:
        vs = Chroma(
            collection_name="compliance_chunks",
            embedding_function=embedding_model,
            persist_directory=CHROMA_DIR
        )
        return vs
    except Exception as e:
        st.warning(f"Failed to load persisted vectorstore: {e}")
        return None

# ---------------------------
# Additional Guardrail helpers (PII, validation, redaction, audit)
# ---------------------------
PII_PATTERNS = {
    "email": re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"),
    "ssn": re.compile(r"\b\d{3}-\d{2}-\d{4}\b"),
    "credit_card": re.compile(r"\b(?:\d[ -]*?){13,16}\b"),
    "phone": re.compile(r"\b(?:\+?\d{1,3}[-.\s]?)?(?:\d{3}[-.\s]?){2}\d{4}\b")
}

def validate_user_query(query: str, max_len: int = 1500):
    if not query or not query.strip():
        return False, "Empty query."
    if len(query) > max_len:
        return False, f"Query too long (>{max_len} chars)."
    banned = ["make a bomb", "how to kill", "bypass security"]
    lowered = query.lower()
    for b in banned:
        if b in lowered:
            return False, "Query contains disallowed content."
    return True, ""

def check_pii_in_text(text: str):
    hits = {}
    if not text:
        return hits
    for k, pat in PII_PATTERNS.items():
        found = pat.findall(text)
        if found:
            hits[k] = found
    return hits

def redact_pii(text: str):
    if not text:
        return text
    out = text
    for pat in PII_PATTERNS.values():
        out = pat.sub("[REDACTED]", out)
    return out

def enforce_output_schema(text: str, max_len: int = 4000):
    if not text:
        return False, "Empty model output."
    if len(text) > max_len:
        return False, "Model output too long."
    refuse_keywords = ["how to make a bomb", "instructions to kill"]
    lower = text.lower()
    for k in refuse_keywords:
        if k in lower:
            return False, "Model output violated safety rules."
    return True, ""

AUDIT_LOG = Path("./audit.log")
def log_audit_entry(event_type: str, payload: dict):
    payload = dict(payload)
    for k, v in payload.items():
        if isinstance(v, str):
            payload[k] = redact_pii(v)
    entry = {"ts": datetime.datetime.utcnow().isoformat() + "Z", "event": event_type, "payload": payload}
    try:
        if AUDIT_LOG.exists():
            AUDIT_LOG.write_text(AUDIT_LOG.read_text() + json.dumps(entry) + "\n")
        else:
            AUDIT_LOG.write_text(json.dumps(entry) + "\n")
    except Exception:
        with AUDIT_LOG.open("a", encoding="utf8") as f:
            f.write(json.dumps(entry) + "\n")

def display_safe_text(text: str):
    safe = redact_pii(text)
    if len(safe) > 10000:
        safe = safe[:10000] + "\n\n...[truncated]"
    st.text_area("Model output (redacted)", value=safe, height=300, key=f"out_{hash(safe)%10000}")

# ---------------------------
# Streamlit UI
# ---------------------------
st.title("CapStone - Enterprise RAG Assistant ")

uploaded_files = st.file_uploader(
    "Upload files (zip, docx, csv, txt, image, pdf)",
    type=["zip", "docx", "csv", "txt", "png", "jpg", "jpeg", "pdf"],
    accept_multiple_files=True
)

if uploaded_files:
    rows = []
    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name
        ext = os.path.splitext(file_name)[1].lower()
        size_mb = uploaded_file.size / (1024 * 1024)
        rows.append([file_name, ext if ext else "No extension", f"{size_mb:.2f} MB"])
    df = pd.DataFrame(rows, columns=["File / Entry", "Format", "Size / Location"])
    st.subheader("Uploaded Files Summary")
    st.table(df)

    if st.button("Extract & Chunk"):
        all_docs: List[Document] = []
        for uploaded in uploaded_files:
            fname = uploaded.name
            raw = uploaded.read()
            extracted = extract_text_from_file(fname, raw)
            all_docs.extend(extracted)

        st.info(f"Extracted {len(all_docs)} raw documents/pages (including image captions).")

        st.session_state["chunks"] = chunk_documents(all_docs)
        chunks = st.session_state.get("chunks", [])
        st.info(f"Created {len(chunks)} chunks (size={CHUNK_SIZE}, overlap={CHUNK_OVERLAP}). Ready to embed.")

        export_df = pd.DataFrame([{
            "source": c.metadata.get("source"),
            "type": c.metadata.get("type", "text"),
            "chunk_index": c.metadata.get("chunk_index"),
            "content": c.page_content
        } for c in chunks])

        st.download_button(
            label="⬇️ Download Chunks as CSV",
            data=export_df.to_csv(index=False).encode("utf-8"),
            file_name="chunks.csv",
            mime="text/csv"
        )

# ---------------------------
# Embedding + ChromaDB
# ---------------------------
if st.button("Embed Chunks to ChromaDB"):
    chunks = st.session_state.get("chunks", [])
    if not chunks:
        st.error("No chunks available. Please upload files and extract chunks first.")
    else:
        st.info("Embedding chunks and storing in ChromaDB...")
        embedding_model = OpenAIEmbeddings(model=EMBEDDING_MODEL)
        vectorstore = Chroma(
            collection_name="compliance_chunks",
            embedding_function=embedding_model,
            persist_directory=CHROMA_DIR
        )

        texts = [chunk.page_content for chunk in chunks]
        metadatas = [chunk.metadata for chunk in chunks]

        try:
            vectorstore.add_texts(texts=texts, metadatas=metadatas)
            vectorstore.persist()
            st.success("✅ Chunks embedded and stored in ChromaDB.")
            st.session_state["vectorstore"] = vectorstore
        except Exception as e:
            st.error(f"Failed to embed/persist chunks: {e}")

# ---------------------------
# Query UI (outside the embed button block)
# ---------------------------
vectorstore = st.session_state.get("vectorstore") or load_persisted_vectorstore()
if vectorstore:
    st.session_state["vectorstore"] = vectorstore  # ensure in session
    mode = st.radio("Choose query mode:", ["Similarity Search", "ChatPrompt QA"])
    query = st.text_input("Enter your query:")

    if query:
        if mode == "Similarity Search":
            try:
                results = vectorstore.similarity_search(query, k=3)
                st.markdown("### 🔎 Similarity Search Results")
                for res in results:
                    st.write(res.page_content)
                    st.write(res.metadata)
            except Exception as e:
                st.error(f"Similarity search failed: {e}")

        elif mode == "ChatPrompt QA":
            ok, msg = validate_user_query(query)
            if not ok:
                st.error(msg)
                log_audit_entry("user_query_blocked", {"query": query, "reason": msg})
																					
            else:
                # input moderation + PII check
                is_flagged, reasons = moderate_text(query)
                pii_hits = check_pii_in_text(query)
                if pii_hits:
                    st.error("Query contains PII and was blocked. Please remove sensitive data.")
                    log_audit_entry("user_query_pii", {"query": query, "pii": pii_hits})
                elif is_flagged:
                    st.error(f"Query blocked by content policy: {', '.join(reasons) or 'unspecified reason'}")
                    log_audit_entry("user_query_moderated", {"query": query, "reasons": reasons})
                else:
																			 
                    try:
                        docs = vectorstore.similarity_search(query, k=3)
                    except Exception as e:
                        st.error(f"Failed to retrieve docs: {e}")
                        log_audit_entry("retrieval_error", {"query": query, "error": str(e)})
                        docs = []

                    # check & redact PII in retrieved context
                    raw_context = "\n\n".join([d.page_content for d in docs])
                    context_pii = check_pii_in_text(raw_context)
                    if context_pii:
                        safe_context = redact_pii(raw_context)
                        st.warning("Retrieved documents contained PII and were redacted before generation.")
                        log_audit_entry("retrieval_pii_redacted", {"query": query, "pii": context_pii})
                    else:
                        safe_context = raw_context

                    # moderation on context
                    ctx_flagged, ctx_reasons = moderate_text(safe_context)
                    if ctx_flagged:
                        st.error("Retrieved documents contain content that violates policy. Remove or redact sensitive content in source files.")
                        log_audit_entry("retrieval_moderated", {"query": query, "reasons": ctx_reasons})
                    else:
																					 
                        system_instructions = (
                            "You are a strict compliance assistant. If a user requests disallowed content "
                            "(illegal instructions, hate, explicit sexual content, or PII extraction), refuse and respond with a short refusal message. "
                            "Always cite sources from the provided context. If the answer is uncertain, say you are unsure and provide best-effort guidance."
                        )

                        user_message = f"Question: {query}\n\nContext:\n{safe_context}"

                        try:
                            client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
                            resp = client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[
                                    {"role": "system", "content": system_instructions},
                                    {"role": "user", "content": user_message}
                                ],
                            )
                            # extract text
                            if isinstance(resp, dict):
                                text_out = resp["choices"][0]["message"]["content"]
                            else:
                                text_out = resp.choices[0].message.content

                            # output enforcement
                            ok_out, reason_out = enforce_output_schema(text_out)
                            if not ok_out:
                                st.error("Model response blocked by output guardrails.")
                                log_audit_entry("output_blocked", {"query": query, "reason": reason_out, "model_text": text_out})
                            else:
                                log_audit_entry("qa_success", {"query": query, "model_text_snippet": text_out[:1000]})
                                st.markdown("### ❓ Question")
                                st.write(query)

                                st.markdown("### 💡 Response")
                                display_safe_text(text_out)

                                st.markdown("### 📑 Sources")
                                for doc in docs:
                                    st.write(doc.metadata)
                        except Exception as e:
                            st.error(f"LLM call failed: {e}")
                            log_audit_entry("llm_error", {"query": query, "error": str(e)})
else:
    st.info("No vectorstore loaded. Upload files and click 'Embed Chunks to ChromaDB' first.")